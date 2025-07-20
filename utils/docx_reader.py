from docx import Document
import re
import logging
from database.database_manager import DatabaseManager
import os


class DocxReader:
    def __init__(self):
        self.db = DatabaseManager()

    # =================================================================
    # CÁC HÀM XỬ LÝ HÌNH ẢNH VÀ HÀM TEST KHÔNG THAY ĐỔI
    # =================================================================
    def extract_images_from_docx(self, file_path):
        """Trích xuất hình ảnh từ file .docx"""
        # Giữ nguyên logic này, nó vẫn có thể hoạt động cho các ảnh được neo trong văn bản
        try:
            doc = Document(file_path)
            images = []
            # ... (Nội dung hàm này được giữ nguyên như cũ) ...
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_data = rel.target_part.blob
                    image_name = os.path.basename(rel.target_ref)
                    images.append({'name': image_name, 'data': image_data,
                                   'paragraph_index': -1})  # paragraph_index không còn quá quan trọng
            logging.info(f"Trích xuất được {len(images)} hình ảnh từ file")
            return images
        except Exception as e:
            logging.error(f"Lỗi trích xuất hình ảnh: {e}")
            return []

    def save_images_to_folder(self, images, output_folder="extracted_images"):
        """Lưu hình ảnh vào thư mục"""
        # Giữ nguyên hàm này
        try:
            if not os.path.exists(output_folder):
                os.makedirs(output_folder)
            saved_images = []
            for i, image in enumerate(images):
                # Đảm bảo tên file là duy nhất nếu cần
                file_path = os.path.join(output_folder, image['name'])
                with open(file_path, 'wb') as f:
                    f.write(image['data'])
                saved_images.append({
                    'original_name': image['name'],
                    'file_path': file_path,
                    'paragraph_index': image['paragraph_index']
                })
                logging.info(f"Đã lưu hình ảnh: {file_path}")
            return saved_images
        except Exception as e:
            logging.error(f"Lỗi lưu hình ảnh: {e}")
            return []

    def extract_image_references_from_text(self, text):
        """Trích xuất tham chiếu hình ảnh từ text"""
        # Giữ nguyên hàm này
        pattern = r'\[file:([^\]]+)\]'
        matches = re.findall(pattern, text)
        return matches

    def process_text_with_images(self, text, images_info):
        """Xử lý text có chứa tham chiếu hình ảnh"""
        # Giữ nguyên hàm này
        image_refs = self.extract_image_references_from_text(text)
        processed_text = text
        for i, ref in enumerate(image_refs):
            image_info = None
            for img in images_info:
                # So sánh tên file gốc được trích xuất
                if os.path.basename(img.get('original_name', '')) == ref:
                    image_info = img
                    break
            if image_info:
                replacement = f"[IMAGE: {image_info['file_path']}]"
                processed_text = processed_text.replace(f"[file:{ref}]", replacement)
            else:
                logging.warning(f"Không tìm thấy hình ảnh cho tham chiếu: {ref}")
        return processed_text

    def test_file_detailed(self, file_path):
        """Test file chi tiết - hiển thị từng dòng và lý do không nhận diện"""
        # LƯU Ý: Hàm này được thiết kế cho định dạng văn bản,
        # cần được viết lại hoàn toàn để hoạt động với định dạng bảng.
        # Tạm thời trả về thông báo.
        return False, "Chức năng test chi tiết chưa được cập nhật cho định dạng bảng."

    # =================================================================
    # PHẦN LOGIC ĐỌC FILE ĐÃ ĐƯỢC CẬP NHẬT HOÀN TOÀN
    # =================================================================
    def read_docx_file(self, file_path, subject_id, creator_id):
        """
        Đọc file .docx và trích xuất câu hỏi từ các BẢNG (TABLES).
        Mỗi câu hỏi được giả định nằm trong một bảng riêng biệt.
        """
        try:
            if not os.path.exists(file_path) or not file_path.lower().endswith('.docx'):
                logging.error(f"File không hợp lệ: {file_path}")
                return False, "File không tồn tại hoặc không phải định dạng .docx"

            logging.info(f"Bắt đầu đọc file theo định dạng BẢNG: {file_path}")
            doc = Document(file_path)
            questions = []

            # Trích xuất và lưu tất cả hình ảnh trong file một lần
            images = self.extract_images_from_docx(file_path)
            saved_images = self.save_images_to_folder(images)
            logging.info(f"Đã trích xuất và lưu {len(saved_images)} hình ảnh")

            # Duyệt qua từng bảng trong tài liệu
            for i, table in enumerate(doc.tables):
                current_question = self._create_empty_question()
                is_valid_question_block = False

                # Duyệt qua từng hàng trong bảng
                for row in table.rows:
                    # Đảm bảo hàng có đủ 2 cột
                    if len(row.cells) < 2:
                        continue

                    label = row.cells[0].text.strip()
                    value = row.cells[1].text.strip()

                    # Nếu ô label trống thì bỏ qua
                    if not label:
                        continue

                    # Bắt đầu một khối câu hỏi khi gặp 'QN=...'
                    if self._is_question_start(label):
                        is_valid_question_block = True
                        current_question['question_number'] = self._extract_question_number(label)
                        # Xử lý tham chiếu hình ảnh ngay tại đây
                        current_question['question_text'] = self.process_text_with_images(value, saved_images)

                    # Chỉ xử lý các dòng khác nếu đã ở trong một khối câu hỏi
                    elif is_valid_question_block:
                        if self._is_option(label):
                            option_letter = label.replace('.', '').strip().upper()
                            current_question['options'][option_letter] = self.process_text_with_images(value,
                                                                                                       saved_images)
                        elif self._is_correct_answer(label):
                            current_question['correct_answer'] = value.strip().upper()
                        elif self._is_mark_info(label):
                            current_question['mark'] = self._extract_mark(value)
                        elif self._is_unit_info(label):
                            current_question['unit'] = value
                        elif self._is_mix_choices_info(label):
                            current_question['mix_choices'] = self._extract_mix_choices(value)

                # Sau khi duyệt xong một bảng, kiểm tra và thêm câu hỏi nếu hợp lệ
                if is_valid_question_block and self._is_valid_question(current_question):
                    questions.append(current_question)
                    logging.info(
                        f"Đã phân tích thành công câu hỏi từ bảng {i + 1} (QN={current_question.get('question_number')})")
                elif is_valid_question_block:
                    logging.warning(
                        f"Bỏ qua câu hỏi không hợp lệ từ bảng {i + 1} (QN={current_question.get('question_number')})")

            logging.info(f"Tổng số câu hỏi hợp lệ đã phân tích được: {len(questions)}")
            if not questions:
                return False, "Không tìm thấy câu hỏi hợp lệ nào trong các bảng của file. Vui lòng kiểm tra định dạng."

            # Lưu vào cơ sở dữ liệu
            saved_count = 0
            for q in questions:
                if self._save_question_to_db(q, subject_id, creator_id):
                    saved_count += 1

            success_message = f"Đã đọc thành công {len(questions)} câu hỏi, lưu {saved_count} câu hỏi vào DB."
            return True, success_message

        except Exception as e:
            logging.error(f"Lỗi nghiêm trọng khi đọc file .docx: {e}", exc_info=True)
            return False, f"Lỗi đọc file: {str(e)}"

    def _create_empty_question(self):
        """Tạo một cấu trúc câu hỏi trống."""
        return {
            'question_text': '',
            'options': {},
            'correct_answer': None,
            'difficulty': 'medium',
            'mark': 1.0,
            'unit': '',
            'mix_choices': False,
            'question_number': None,
        }

    # =================================================================
    # CÁC HÀM HELPER ĐÃ ĐƯỢC CẬP NHẬT
    # =================================================================
    def _is_question_start(self, text):
        """Kiểm tra label ở cột 1 có phải là bắt đầu câu hỏi không"""
        return re.match(r'^QN\s*=\s*\d+', text, re.IGNORECASE)

    def _extract_question_number(self, text):
        """Trích xuất số câu hỏi từ label cột 1"""
        match = re.search(r'(\d+)', text)
        return int(match.group(1)) if match else None

    def _is_option(self, text):
        """Kiểm tra label ở cột 1 có phải là đáp án không"""
        return re.match(r'^[a-d]\.', text, re.IGNORECASE)

    def _is_correct_answer(self, text):
        """Kiểm tra label ở cột 1 có phải là đáp án đúng không"""
        return text.strip().upper().startswith('ANSWER')

    def _is_mark_info(self, text):
        """Kiểm tra label ở cột 1 có phải là thông tin điểm không"""
        return text.strip().upper().startswith('MARK')

    def _is_unit_info(self, text):
        """Kiểm tra label ở cột 1 có phải là thông tin đơn vị không"""
        return text.strip().upper().startswith('UNIT')

    def _is_mix_choices_info(self, text):
        """Kiểm tra label ở cột 1 có phải là thông tin trộn đáp án không"""
        return text.strip().upper().startswith('MIX CHOICES')

    def _extract_option(self, text):
        # Hàm này không còn cần thiết trong logic đọc bảng
        pass

    def _extract_correct_answer(self, text):
        # Không còn cần thiết, vì giá trị được lấy trực tiếp từ cột 2
        pass

    def _extract_mark(self, value_text):
        """Trích xuất điểm số từ value ở cột 2"""
        match = re.search(r'(\d+\.?\d*)', value_text)
        try:
            return float(match.group(1)) if match else 1.0
        except (ValueError, AttributeError):
            return 1.0

    def _extract_unit(self, text):
        # Không còn cần thiết, giá trị được lấy trực tiếp
        pass

    def _extract_mix_choices(self, value_text):
        """Trích xuất thông tin trộn đáp án từ value ở cột 2"""
        return value_text.lower().strip() in ['yes', 'true', 'có']

    def _is_valid_question(self, question):
        """Kiểm tra câu hỏi có hợp lệ không"""
        if not question:
            return False
        # Kiểm tra các điều kiện cơ bản
        has_text = bool(question.get('question_text'))
        has_enough_options = len(question.get('options', {})) >= 4
        has_correct_answer_format = question.get('correct_answer') in ['A', 'B', 'C', 'D']
        # Kiểm tra xem đáp án đúng có nằm trong các lựa chọn không
        correct_answer_in_options = question.get('correct_answer') in question.get('options', {})

        return has_text and has_enough_options and has_correct_answer_format and correct_answer_in_options

    def _save_question_to_db(self, question, subject_id, creator_id):
        """Lưu câu hỏi vào cơ sở dữ liệu"""
        try:
            # Logic lưu DB được giữ nguyên
            additional_info = []
            if question.get('unit'):
                additional_info.append(f"Unit: {question['unit']}")
            if question.get('mark') and question['mark'] != 1.0:
                additional_info.append(f"Mark: {question['mark']}")
            if question.get('mix_choices'):
                additional_info.append("Mix Choices: Yes")

            full_question_text = question['question_text']
            if additional_info:
                full_question_text += f"\n[{' | '.join(additional_info)}]"

            query = """
                    INSERT INTO questions (subject_id, question_text, option_a, option_b, option_c, option_d, \
                                           correct_answer, difficulty_level, created_by)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s) \
                    """
            params = (
                subject_id,
                full_question_text,
                question['options'].get('A', ''),
                question['options'].get('B', ''),
                question['options'].get('C', ''),
                question['options'].get('D', ''),
                question['correct_answer'],
                question['difficulty'],
                creator_id
            )
            self.db.execute_query(query, params)
            return True
        except Exception as e:
            logging.error(f"Lỗi lưu câu hỏi (QN={question.get('question_number')}) vào DB: {e}")
            return False

    def get_template_instructions(self):
        """Trả về hướng dẫn định dạng template"""
        # Cập nhật hướng dẫn cho định dạng bảng
        return """
HƯỚNG DẪN ĐỊNH DẠNG FILE .DOCX (DẠNG BẢNG)

Mỗi câu hỏi phải nằm trong một bảng (Table) riêng biệt. Bảng phải có 2 cột.

- Cột 1: Chứa các nhãn (Label) như 'QN=1', 'a.', 'ANSWER:', 'MARK:'.
- Cột 2: Chứa nội dung tương ứng.

Cấu trúc một bảng câu hỏi mẫu:
|------------------|----------------------------------------------------|
|   **Cột 1 (Nhãn)** |   **Cột 2 (Nội dung)** |
|------------------|----------------------------------------------------|
| QN=1             | Nội dung câu hỏi [file:image.jpg]                  |
| a.               | Nội dung đáp án A                                  |
| b.               | Nội dung đáp án B                                  |
| c.               | Nội dung đáp án C                                  |
| d.               | Nội dung đáp án D                                  |
| ANSWER:          | B                                                  |
| MARK:            | 0.5                                                |
| UNIT:            | Chapter1                                           |
| MIX CHOICES:     | Yes                                                |
|------------------|----------------------------------------------------|

LƯU Ý:
- Mỗi câu hỏi phải là một bảng riêng.
- Các nhãn ở Cột 1 phải chính xác (ví dụ: 'QN=1', 'a.', 'ANSWER:').
- Hình ảnh sẽ được tự động trích xuất và liên kết nếu có tham chiếu [file:...].
"""
